from googlesearch import search
from bs4 import BeautifulSoup
import requests
from docx import Document
from docx.shared import Inches

# Search query
query = "Eryngium billardieri"

# Number of search results
num_results = 40

# Perform the search
search_results = search(query, num_results=num_results, lang='en')

# Create a Word document
doc = Document()
doc.add_heading(f"Articles about '{query}'", level=1)

# Create a table with headers
table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'
table.autofit = False
table.columns[0].width = Inches(0.5)
table.columns[1].width = Inches(2)
table.columns[2].width = Inches(2)
table.columns[3].width = Inches(4)

# Set the header text
header_cells = table.rows[0].cells
header_cells[0].text = "No."
header_cells[1].text = "Article"
header_cells[2].text = "Link"
header_cells[3].text = "Abstract"

# Add the search results to the table
for index, result in enumerate(search_results):
    row_cells = table.add_row().cells
    row_cells[0].text = str(index + 1)
    row_cells[2].text = result

    try:
        response = requests.get(result)
        soup = BeautifulSoup(response.text, 'html.parser')
        title = soup.title.text.strip()
        abstract = soup.find('meta', attrs={'name': 'description'})
        abstract = abstract['content'].strip() if abstract else "N/A"
        row_cells[1].text = title
        row_cells[3].text = abstract
    except:
        row_cells[1].text = "N/A"
        row_cells[3].text = "N/A"

# Save the document
doc.save("articles.docx")

print(f"Articles about '{query}' saved to 'articles.docx'.")
